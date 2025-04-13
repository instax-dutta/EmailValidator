import re
import os
import csv
import dns.resolver
import dns.exception
import logging
import time
import smtplib
import socket
from concurrent.futures import ThreadPoolExecutor, as_completed
import openpyxl
import xlrd
import argparse
from email.utils import parseaddr
from functools import lru_cache
from ratelimit import limits, sleep_and_retry
from cachetools import cached, TTLCache
from idna import encode as idna_encode
from docx import Document

# Cache configurations
dns_cache = TTLCache(maxsize=1000, ttl=3600)  # 1 hour TTL
smtp_cache = TTLCache(maxsize=1000, ttl=1800)  # 30 minutes TTL

# Rate limiting configurations
MAX_CALLS_PER_SECOND = 5
SMTP_CALLS_PER_MINUTE = 30

# Constants
MAX_WORKERS = 20  # Concurrent validation threads
DNS_TIMEOUT = 5   # Seconds for DNS timeout
SMTP_TIMEOUT = 10  # Seconds for SMTP timeout
LOG_FILE = "email_validation.log"
VALID_EMAIL_FILE = "valid_emails.txt"

# More comprehensive regex for email validation
EMAIL_REGEX = r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"

# List of common disposable email domains
DISPOSABLE_DOMAINS = [
    'mailinator.com', 'yopmail.com', 'tempmail.com', 'temp-mail.org', 
    'guerrillamail.com', '10minutemail.com', 'trashmail.com', 'sharklasers.com',
    'throwawaymail.com', 'getairmail.com', 'mailnesia.com', 'tempinbox.com',
    'dispostable.com', 'mailcatch.com', 'anonbox.net', 'getnada.com'
]

# Common role-based email prefixes
ROLE_BASED_PREFIXES = [
    'admin', 'info', 'support', 'sales', 'contact', 'help', 'webmaster',
    'postmaster', 'hostmaster', 'abuse', 'noreply', 'no-reply', 'mail',
    'office', 'marketing', 'team', 'billing', 'jobs', 'career', 'hr'
]

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler()
    ]
)

def is_syntactically_valid(email):
    """Checks if the email has a valid basic format."""
    # Use parseaddr to handle edge cases better
    _, addr = parseaddr(email)
    if not addr or addr != email:
        return False
    
    # Check basic pattern
    if not re.match(EMAIL_REGEX, email):
        return False
    
    # Additional syntax checks
    if '..' in email:  # No consecutive dots
        return False
    
    # Check for valid domain part
    try:
        username, domain = email.split('@')
        if not username or not domain:
            return False
        if domain.startswith('.') or domain.endswith('.'):
            return False
        if not re.match(r'^[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', domain):
            return False
    except ValueError:
        return False
    
    return True

@cached(cache=dns_cache)
def check_mx_record(domain):
    """Checks if the domain has a valid MX record with caching."""
    try:
        resolver.resolve(domain, 'MX')
        logging.debug(f"MX record found for domain: {domain}")
        return True
    except dns.resolver.NoAnswer:
        logging.warning(f"MX check failed for {domain}: No MX record found (NoAnswer).")
        return False
    except dns.resolver.NXDOMAIN:
        logging.warning(f"MX check failed for {domain}: Domain does not exist (NXDOMAIN).")
        return False
    except dns.resolver.Timeout:
        logging.error(f"MX check failed for {domain}: DNS query timed out after {DNS_TIMEOUT} seconds.")
        return False # Treat timeout as failure for robustness
    except dns.exception.DNSException as e:
        logging.error(f"MX check failed for {domain}: DNS query error: {e}")
        return False
    except Exception as e:
        logging.error(f"Unexpected error during MX check for {domain}: {e}")
        return False

def is_disposable_email(domain):
    """Check if the email domain is a known disposable email service."""
    return domain.lower() in DISPOSABLE_DOMAINS

def is_role_based_email(username):
    """Check if the email is a role-based address."""
    return username.lower() in ROLE_BASED_PREFIXES

@sleep_and_retry
@limits(calls=SMTP_CALLS_PER_MINUTE, period=60)
@cached(cache=smtp_cache)
def verify_smtp_connection(email, domain):
    """Attempt to verify email by connecting to the SMTP server with rate limiting and caching."""
    try:
        # Get MX records and sort by preference
        mx_records = sorted([(pref, str(exch)) for pref, exch in dns.resolver.resolve(domain, 'MX')])
        if not mx_records:
            return False
        
        # Connection pool configuration
        smtp_pool = {}
        
        # Try to connect to the mail server
        for _, mail_server in mx_records:
            try:
                if mail_server in smtp_pool:
                    smtp = smtp_pool[mail_server]
                else:
                    smtp = smtplib.SMTP(timeout=SMTP_TIMEOUT)
                    smtp.connect(mail_server)
                    smtp.helo(socket.getfqdn())
                    smtp_pool[mail_server] = smtp
                
                smtp.mail('')
                code, _ = smtp.rcpt(email)
                
                # 250 is the success code for RCPT command
                if code == 250:
                    return True
                # Some servers don't allow email verification but the domain is still valid
                if code in [421, 450, 451, 452]:
                    return True
            except (socket.timeout, smtplib.SMTPServerDisconnected, 
                    smtplib.SMTPConnectError, ConnectionRefusedError):
                continue
            except Exception as e:
                logging.debug(f"SMTP verification error for {email} on server {mail_server}: {e}")
                # If we get a specific error, the server exists but rejects our probe
                # This is still a potentially valid email
                return True
        
        return False
    except Exception as e:
        logging.debug(f"SMTP verification failed for {email}: {e}")
        # If MX records exist but SMTP verification fails, we'll still consider it potentially valid
        return True

def validate_email(email):
    """Performs comprehensive email validation with IDN support."""
    if not isinstance(email, str) or not email.strip():
        return None # Skip empty or non-string entries

    email = email.strip()
    
    # Handle IDN domains
    try:
        if '@' in email:
            local_part, domain = email.split('@')
            if any(ord(c) > 127 for c in domain):  # Contains non-ASCII
                domain = idna_encode(domain).decode('ascii')
                email = f"{local_part}@{domain}"
    except Exception as e:
        logging.debug(f"IDN conversion failed for {email}: {e}")
        return None
    if not is_syntactically_valid(email):
        logging.debug(f"Invalid syntax: {email}")
        return None

    try:
        username, domain = email.split('@')
    except IndexError:
        logging.warning(f"Could not split email into user/domain: {email}")
        return None

    # Check for disposable email domains
    if is_disposable_email(domain):
        logging.debug(f"Disposable email detected: {email}")
        return None

    # Check for role-based emails
    if is_role_based_email(username):
        logging.debug(f"Role-based email detected: {email}")
        return None

    # Check MX record
    if not check_mx_record(domain):
        # Logging is handled within check_mx_record
        return None
    
    # Optional: Verify SMTP connection (can be disabled for performance)
    # Uncomment the following lines to enable SMTP verification
    # if not verify_smtp_connection(email, domain):
    #     logging.debug(f"SMTP verification failed: {email}")
    #     return None

    logging.debug(f"Validated email: {email}")
    return email

def read_txt(filepath):
    """Reads emails line-by-line from a .txt file."""
    emails = set()
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            for i, line in enumerate(f):
                try:
                    # Find all potential emails within a line
                    found = re.findall(EMAIL_REGEX, line)
                    for email in found:
                        emails.add(email.strip())
                except Exception as line_error:
                    logging.warning(f"Error processing line {i+1} in {filepath}: {line_error} - Line: '{line.strip()}'")
    except FileNotFoundError:
        logging.error(f"Error: File not found at {filepath}")
        return [] # Return empty list on critical error
    except Exception as e:
        logging.error(f"Critical error reading TXT file {filepath}: {e}")
        return [] # Return empty list on critical error
    logging.info(f"Extracted {len(emails)} unique potential emails from {filepath}")
    return list(emails)

def read_csv(filepath):
    """Reads emails line-by-line from a .csv file, checking all cells."""
    emails = set()
    try:
        with open(filepath, 'r', newline='', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for i, row in enumerate(reader):
                try:
                    for cell in row:
                        if isinstance(cell, str):
                            # Find all potential emails within a cell
                            found = re.findall(EMAIL_REGEX, cell)
                            for email in found:
                                emails.add(email.strip())
                except Exception as row_error:
                    logging.warning(f"Error processing row {i+1} in {filepath}: {row_error} - Row: '{row}'")
    except FileNotFoundError:
        logging.error(f"Error: File not found at {filepath}")
        return []
    except csv.Error as e:
        logging.error(f"CSV formatting error in file {filepath}, line {reader.line_num}: {e}")
        # Attempt to continue if possible, or return [] if critical
        return []
    except Exception as e:
        logging.error(f"Critical error reading CSV file {filepath}: {e}")
        return []
    logging.info(f"Extracted {len(emails)} unique potential emails from {filepath}")
    return list(emails)

def read_excel(filepath):
    # Note: Reading large Excel files can still be memory intensive with these libraries.
    # For truly massive Excel files (> millions of rows), specialized libraries or
    # converting to CSV first might be necessary.
    """Reads emails from an .xlsx or .xls file, checking all cells."""
    emails = set()
    try:
        if filepath.endswith('.xlsx'):
            workbook = openpyxl.load_workbook(filepath, read_only=True, data_only=True) # data_only=True tries to get cell values instead of formulas
            logging.info(f"Reading sheets from {filepath}...")
            for sheet_name in workbook.sheetnames:
                logging.debug(f"Processing sheet: {sheet_name}")
                sheet = workbook[sheet_name]
                # Using iter_rows for potentially better memory usage than direct indexing for large sheets
                for row_idx, row in enumerate(sheet.iter_rows()):
                    try:
                        for cell in row:
                            if cell.value and isinstance(cell.value, str):
                                found = re.findall(EMAIL_REGEX, cell.value)
                                for email in found:
                                    emails.add(email.strip())
                    except Exception as cell_error:
                         logging.warning(f"Error processing cell in sheet '{sheet_name}', row {row_idx+1} of {filepath}: {cell_error}")

        elif filepath.endswith('.xls'):
             # xlrd handles memory reasonably well for .xls
            workbook = xlrd.open_workbook(filepath, logfile=open(os.devnull, 'w')) # Suppress xlrd warnings
            logging.info(f"Reading sheets from {filepath}...")
            for sheet in workbook.sheets():
                logging.debug(f"Processing sheet: {sheet.name}")
                try:
                    for row_idx in range(sheet.nrows):
                        for col_idx in range(sheet.ncols):
                            cell_value = sheet.cell_value(row_idx, col_idx)
                            if isinstance(cell_value, str):
                                found = re.findall(EMAIL_REGEX, cell_value)
                                for email in found:
                                    emails.add(email.strip())
                except Exception as sheet_error:
                    logging.warning(f"Error processing sheet '{sheet.name}' in {filepath}: {sheet_error}")
        else:
             logging.error(f"Unsupported Excel file format: {filepath}")
             return []

    except FileNotFoundError:
        logging.error(f"Error: File not found at {filepath}")
        return []
    except Exception as e:
        # Catch potential errors from openpyxl/xlrd loading
        logging.error(f"Critical error reading Excel file {filepath}: {e}")
        return []
    logging.info(f"Extracted {len(emails)} unique potential emails from {filepath}")
    return list(emails)

def read_doc(filepath):
    """Reads emails from a .doc or .docx file."""
    emails = set()
    try:
        doc = Document(filepath)
        for paragraph in doc.paragraphs:
            if isinstance(paragraph.text, str):
                found = re.findall(EMAIL_REGEX, paragraph.text)
                for email in found:
                    emails.add(email.strip())
        
        # Also check tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if isinstance(cell.text, str):
                        found = re.findall(EMAIL_REGEX, cell.text)
                        for email in found:
                            emails.add(email.strip())
    except Exception as e:
        logging.error(f"Error reading DOC file {filepath}: {e}")
        return []
    
    logging.info(f"Extracted {len(emails)} unique potential emails from {filepath}")
    return list(emails)

def read_emails_from_file(filepath):
    """Determines file type and calls the appropriate reader function."""
    _, extension = os.path.splitext(filepath)
    extension = extension.lower()

    logging.info(f"Reading file: {filepath} (Extension: {extension})")

    if extension == '.txt':
        return read_txt(filepath)
    elif extension == '.csv':
        return read_csv(filepath)
    elif extension in ['.xlsx', '.xls']:
        return read_excel(filepath)
    elif extension in ['.doc', '.docx']:
        return read_doc(filepath)
    else:
        logging.error(f"Unsupported file type: {extension}. Please provide a .txt, .csv, .xlsx, .xls, .doc, or .docx file.")
        return []

def main(args): # Accept args object
    input_filepath = args.input_file

    if not os.path.exists(input_filepath):
        logging.error(f"Input file not found: {input_filepath}")
        return

    logging.info(f"Starting email validation process for: {input_filepath}")
    emails_to_validate = read_emails_from_file(input_filepath)

    if not emails_to_validate:
        logging.warning("No emails found or extracted from the input file.")
        return

    logging.info(f"Found {len(emails_to_validate)} potential emails to validate.")

    valid_emails = []
    processed_count = 0
    start_time = time.time() # Start timer

    # Use ThreadPoolExecutor for concurrent DNS checks
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        # Using set to avoid submitting duplicate emails for validation
        unique_emails = set(emails_to_validate)
        logging.info(f"Submitting {len(unique_emails)} unique emails for validation...")
        future_to_email = {executor.submit(validate_email, email): email for email in unique_emails}
        total_emails = len(future_to_email)

        for future in as_completed(future_to_email):
            email = future_to_email[future]
            processed_count += 1
            try:
                result = future.result()
                if result:
                    valid_emails.append(result)
                # Log progress periodically
                if processed_count % 100 == 0 or processed_count == total_emails:
                     # Calculate elapsed time and estimate remaining time
                     elapsed_time = time.time() - start_time
                     emails_per_second = processed_count / elapsed_time if elapsed_time > 0 else 0
                     remaining_emails = total_emails - processed_count
                     estimated_time_remaining = (remaining_emails / emails_per_second) if emails_per_second > 0 else 0
                     logging.info(f"Processed {processed_count}/{total_emails} emails... ({emails_per_second:.2f} emails/sec, ETA: {estimated_time_remaining:.1f}s)")

            except Exception as exc:
                # Log exceptions raised during the validate_email task itself
                logging.error(f'Validation task for {email} generated an exception: {exc}')

    end_time = time.time() # End timer
    total_time = end_time - start_time
    logging.info(f"Validation complete. Found {len(valid_emails)} valid emails out of {len(unique_emails)} unique emails processed.")
    logging.info(f"Total processing time: {total_time:.2f} seconds.")

    if valid_emails:
        # Sort emails before writing
        valid_emails.sort()
        try:
            with open(VALID_EMAIL_FILE, 'w', encoding='utf-8') as f:
                for email in valid_emails: # Already unique from processing set
                    f.write(email + '\n')
            logging.info(f"Valid emails saved to: {VALID_EMAIL_FILE}")
        except IOError as e:
            logging.error(f"Error writing valid emails to file {VALID_EMAIL_FILE}: {e}")
        except Exception as e:
            logging.error(f"Unexpected error writing output file: {e}")
    else:
        logging.warning("No valid emails found to write to the output file.")

    logging.info(f"Process finished. Check '{LOG_FILE}' for detailed logs.")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Validate emails from a file.')
    parser.add_argument('input_file', help='Path to the input file containing emails (.txt, .csv, .xlsx, or .xls)')
    args = parser.parse_args()
    main(args)
